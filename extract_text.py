from docx import Document
import os

# Lista de archivos DOCX
docx_files = [
    'Material estudio radio/Configuraciones y Capacidades de Radios Ericsson Minilink - 2024.docx',
    'Material estudio radio/manual_introduccion_redes_tcp_ip.docx',
    'Material estudio radio/MW Radio Propagation.docx'
]

# Extraer texto de cada archivo
for docx_file in docx_files:
    try:
        print(f'\nExtrayendo texto de: {os.path.basename(docx_file)}...')
        
        # Abrir documento
        doc = Document(docx_file)
        
        # Crear archivo de texto
        txt_file = docx_file.replace('.docx', '.txt')
        
        # Extraer todo el texto
        with open(txt_file, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                if para.text.strip():  # Solo escribir párrafos no vacíos
                    f.write(para.text + '\n')
            
            # También extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            f.write(cell.text + '\n')
        
        print(f'✓ Texto extraído a: {os.path.basename(txt_file)}')
    except Exception as e:
        print(f'✗ Error: {str(e)}')

print('\n¡Extracción completada!')
